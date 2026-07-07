import os
import requests
import re
import base64  # 🔥 Importación necesaria para convertir el audio a string seguro
import json    # 🌟 Necesario para procesar el string JSON de la cuenta de servicio
from bs4 import BeautifulSoup
from docx import Document as DocxDocument  # Evita colisión de nombres
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi import BackgroundTasks
from pydantic import BaseModel
import chromadb
# Usamos el protocolo base de Chroma para registrar funciones de embedding personalizadas
from chromadb.api.types import EmbeddingFunction, Documents, Embeddings
from google import genai
from google.genai import types
from google.cloud import texttospeech  # 🔥 Importación oficial de Google Cloud TTS
from langchain_core.documents import Document as LangchainDocument  # Estructura de datos base
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter


# 1. CONFIGURACIÓN DEL SERVIDOR
app = FastAPI(title="Cerebro Local/Híbrido del Asistente Universitario")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

API_KEY = os.getenv("GEMINI_API_KEY")

if not API_KEY:
    raise RuntimeError("No existe GEMINI_API_KEY")

cliente_gemini = genai.Client(api_key=API_KEY)


# =====================================================================
# 🔥 CONFIGURACIÓN SEGURA DE GOOGLE CLOUD TTS PARA RENDER
# =====================================================================
cliente_tts = None

# 1. Intentamos leer las credenciales en formato JSON de texto directo desde la variable
google_json_str = os.getenv("GOOGLE_CREDENTIALS_JSON")

if google_json_str:
    try:
        # Convertimos la cadena de texto de la variable en un diccionario de Python
        credenciales_dict = json.loads(google_json_str)
        # Inicializamos el cliente pasándole directamente las credenciales estructuradas
        cliente_tts = texttospeech.TextToSpeechClient.from_service_account_info(credenciales_dict)
        print("✅ Cliente de Google Cloud TTS inicializado correctamente desde variable de entorno (JSON).")
    except Exception as e:
        print(f"❌ Error al parsear GOOGLE_CREDENTIALS_JSON: {e}")
else:
    # 2. Respaldo por si estás probando localmente con el archivo .json tradicional
    ruta_archivo = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
    if ruta_archivo and os.path.exists(ruta_archivo):
        cliente_tts = texttospeech.TextToSpeechClient()
        print(f"✅ Cliente TTS inicializado desde archivo local: {ruta_archivo}")

if not cliente_tts:
    print("⚠️ ADVERTENCIA: No se configuraron credenciales para Google Cloud TTS. El endpoint /preguntar fallará al generar voz.")
# =====================================================================


CARPETA_DOCUMENTOS = "./documentos_uni"
CARPETA_DB_VECTORIAL = "./db_vectorial"
ARCHIVO_URLS = "./urls_web.txt"

CACHE_RESPUESTAS = {}

if not os.path.exists(CARPETA_DOCUMENTOS):
    os.makedirs(CARPETA_DOCUMENTOS)

# Inicializar ChromaDB Local
cliente_chroma = chromadb.PersistentClient(path=CARPETA_DB_VECTORIAL)

# 🔥 CLASE DE EMBEDDING REPARADA
class GeminiEmbeddingFunctionCloud(EmbeddingFunction):
    def __call__(self, input: Documents) -> Embeddings:
        try:
            # Llamamos al modelo usando el identificador limpio compatible con el SDK v1/v1beta unificado
            respuesta = cliente_gemini.models.embed_content(
                model="gemini-embedding-001",
                contents=list(input)
            )
            
            # Extraemos los vectores numéricos devueltos por el cliente
            return [e.values for e in respuesta.embeddings]
            
        except Exception as e:
            print(f"❌ Error al generar embeddings con el SDK oficial de Google: {e}")
            raise e

# Instanciamos nuestra función cloud optimizada mediante HTTP requests
funcion_embedding_cloud = GeminiEmbeddingFunctionCloud()

# Creamos o cargamos la colección vinculada a la nueva función integrada
coleccion = cliente_chroma.get_or_create_collection(
    name="universidad_docs",
    embedding_function=funcion_embedding_cloud
)

# --- FUNCIONES AUXILIARES DE EXTRACCIÓN Y LECTURA ---

def obtener_urls_desde_archivo():
    """Lee las URLs desde el archivo de texto urls_web.txt línea por línea."""
    if not os.path.exists(ARCHIVO_URLS):
        with open(ARCHIVO_URLS, "w", encoding="utf-8") as f:
            f.write("# Coloca aquí una URL por línea (ejemplo: https://www.google.com)\n")
        return []
    
    urls = []
    with open(ARCHIVO_URLS, "r", encoding="utf-8") as f:
        for linea in f:
            linea_limpia = linea.strip()
            if linea_limpia and not linea_limpia.startswith("#"):
                urls.append(linea_limpia)
    return urls

def extraer_texto_de_word(ruta_archivo):
    """Extrae todo el texto de un archivo .docx manteniendo saltos de línea básicos."""
    try:
        doc = DocxDocument(ruta_archivo)
        texto_completo = []
        for parrafo in doc.paragraphs:
            if parrafo.text.strip():
                texto_completo.append(parrafo.text)
        return "\n".join(texto_completo)
    except Exception as e:
        print(f"⚠️ Error al leer el archivo Word {ruta_archivo}: {e}")
        return None

def extraer_texto_de_url(url):
    """Descarga una página web y extrae el texto sumamente limpio y normalizado."""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        respuesta = requests.get(url, headers=headers, timeout=10)
        if respuesta.status_code == 200:
            respuesta.encoding = 'utf-8'
            soup = BeautifulSoup(respuesta.text, 'html.parser')
            
            for elemento in soup(["script", "style", "nav", "footer", "header", "aside", "form"]):
                elemento.decompose()
                
            texto_sucio = soup.get_text(separator=" ")
            texto_limpio = re.sub(r'\s+', ' ', texto_sucio).strip()
            return texto_limpio
    except Exception as e:
        print(f"⚠️ Error al raspar la URL {url}: {e}")
    return None

# 2. PROCESAMIENTO MULTIFUENTE E INCREMENTAL
def cargar_y_vectorizar_fuentes():
    print("🔍 Verificando fuentes en la base de datos vectorial...")
    
    fuentes_existentes = set()
    if coleccion.count() > 0:
        resultados = coleccion.get(include=["metadatas"])
        for meta in resultados.get("metadatas", []):
            if meta and "fuente" in meta:
                fuentes_existentes.add(meta["fuente"])

    todos_los_documentos_nuevos = []

    # Procesar archivos locales
    if os.path.exists(CARPETA_DOCUMENTOS):
        for archivo in os.listdir(CARPETA_DOCUMENTOS):
            ruta_completa = os.path.join(CARPETA_DOCUMENTOS, archivo)
            
            # 🔥 Omitimos los archivos temporales de Word para evitar fallas
            if archivo.startswith("~$"):
                continue
                
            if archivo in fuentes_existentes:
                continue
                
            if archivo.endswith('.pdf'):
                print(f"🆕 Detectado nuevo PDF: {archivo}")
                try:
                    loader = PyPDFLoader(ruta_completa)
                    todos_los_documentos_nuevos.extend(loader.load())
                except Exception as e:
                    print(f"⚠️ Error al cargar el PDF {archivo}: {e}")
                
            elif archivo.endswith('.docx'):
                print(f"🆕 Detectado nuevo Word: {archivo}")
                texto_word = extraer_texto_de_word(ruta_completa)
                if texto_word:
                    todos_los_documentos_nuevos.append(
                        LangchainDocument(page_content=texto_word, metadata={"fuente": archivo})
                    )

    # Cargar URLs
    paginas_web_objetivo = obtener_urls_desde_archivo()

    for url in paginas_web_objetivo:
        if url in fuentes_existentes:
            continue
        print(f"🆕 Detectada nueva URL desde archivo txt: {url}")
        texto_web = extraer_texto_de_url(url)
        if texto_web:
            todos_los_documentos_nuevos.append(
                LangchainDocument(page_content=texto_web, metadata={"fuente": url})
            )

    if not todos_los_documentos_nuevos:
        print(f"✅ Base de datos al día. No hay fuentes nuevas. Total en DB: {coleccion.count()} fragmentos.")
        return

    print(f"Procesando {len(todos_los_documentos_nuevos)} nuevas fuentes de información...")
    
    splitter = RecursiveCharacterTextSplitter(chunk_size=750, chunk_overlap=130)
    fragmentos = splitter.split_documents(todos_los_documentos_nuevos)
    
    textos = [frag.page_content for frag in fragmentos]
    metadatos = [frag.metadata if frag.metadata else {"fuente": "desconocida"} for frag in fragmentos]
    
    ids = []
    for i, frag in enumerate(fragmentos):
        nombre_fuente = frag.metadata.get("fuente", "desconocida")
        nombre_limpio = "".join(c for c in nombre_fuente if c.isalnum() or c in "._-")
        ids.append(f"id_{nombre_limpio}_chunk_{i}")
    
    # Subida por lotes (Batching) para evitar el error 400 de Google
    TAMANO_LOTE = 90
    total_fragmentos = len(fragmentos)
    print(f"📦 Dividiendo {total_fragmentos} fragmentos en lotes de {TAMANO_LOTE} para la API de Google...")

    for inicio in range(0, total_fragmentos, TAMANO_LOTE):
        fin = min(inicio + TAMANO_LOTE, total_fragmentos)
        
        lote_textos = textos[inicio:fin]
        lote_metadatos = metadatos[inicio:fin]
        lote_ids = ids[inicio:fin]
        
        print(f"🚀 Enviando lote: fragmentos del {inicio} al {fin}...")
        coleccion.add(
            documents=lote_textos,
            metadatas=lote_metadatos,
            ids=lote_ids
        )
        
    print(f"✅ ¡Éxito! Se añadieron/actualizaron todos los {total_fragmentos} fragmentos. Total acumulado en DB: {coleccion.count()}.")

import asyncio

@app.on_event("startup")
async def startup_event():
    print("🚀 Servidor en línea. Iniciando verificación de archivos en segundo plano de forma segura...")
    asyncio.create_task(asyncio.to_thread(cargar_y_vectorizar_fuentes))


# 3. ENDPOINTS DEL SERVIDOR

@app.get("/")
def ruta_raiz():
    return {"status": "online", "mensaje": "Cerebro del Asistente Universitario Activo"}


class Consulta(BaseModel):
    pregunta: str

@app.post("/preguntar")
async def responder_pregunta(consulta: Consulta):
    try:
        pregunta_normalizada = consulta.pregunta.strip().lower()
        
        # Guardaremos el texto de la respuesta final aquí
        texto_final = ""

        # Verificamos si la respuesta ya existe en caché para mitigar consumo de tokens
        if pregunta_normalizada in CACHE_RESPUESTAS:
            print("🚀 Respuesta entregada desde la caché local (0 llamadas consumidas a Gemini)")
            texto_final = CACHE_RESPUESTAS[pregunta_normalizada]
        else:
            contexto = ""
            
            if coleccion.count() > 0:
                resultado_busqueda = coleccion.query(
                    query_texts=[consulta.pregunta],
                    n_results=3
                )
                if resultado_busqueda and 'documents' in resultado_busqueda and resultado_busqueda['documents']:
                    documentos_encontrados = resultado_busqueda['documents'][0]
                    documentos_limpios = [doc for doc in documentos_encontrados if doc]
                    contexto = "\n".join(documentos_limpios)

            prompt_sistema = (
                "Eres el asistente virtual interactivo oficial de la universidad cesar vallejo de la sede o campus de la ciudad de Tarapoto, representado por un avatar en una pantalla.\n"
                "Tu objetivo es ayudar amablemente a estudiantes y visitantes con información del campus.\n"
                "Usa exclusivamente el siguiente contexto de la documentación institucional para responder la pregunta.\n"
                "Si no sabes la respuesta o no se encuentra en el contexto, di amablemente: 'Lo siento, solo doy información que se encuentra en mis registros institucionales'. No inventes datos.\n"
                "Sé claro y amable (máximo 4 oraciones), ya que tu respuesta será leída en un monitor público.\n"
                "evita decir Hola en cada respuesta.\n\n"
                f"Contexto Institucional:\n{contexto}"
            )

            respuesta = cliente_gemini.models.generate_content(
                model="gemini-2.5-flash",
                contents=consulta.pregunta,
                config=types.GenerateContentConfig(
                    system_instruction=prompt_sistema,
                    temperature=0.3
                )
            )
            texto_final = respuesta.text
            CACHE_RESPUESTAS[pregunta_normalizada] = texto_final

        # Validamos que el cliente de voz esté inicializado correctamente
        if not cliente_tts:
            raise HTTPException(status_code=500, detail="El cliente de Google Cloud TTS no está configurado.")

        # 🔥 SÍNTESIS DE AUDIO CON GOOGLE CLOUD TEXT-TO-SPEECH
        # Configuramos el texto de entrada que el motor de voz va a procesar
        synthesis_input = texttospeech.SynthesisInput(text=texto_final)
        
        # Seleccionamos parámetros de una voz premium en español (ejemplo: es-US-Neural2-B)
        voice = texttospeech.VoiceSelectionParams(
            language_code="es-US",
            name="es-US-Neural2-B" 
        )
        
        # Definimos el formato MP3 de salida y pequeños ajustes naturales
        audio_config = texttospeech.AudioConfig(
            audio_encoding=texttospeech.AudioEncoding.MP3,
            speaking_rate=1.05,
            pitch=0.0
        )
        
        # Enviamos la solicitud de generación de voz a Google Cloud
        response_audio = cliente_tts.synthesize_speech(
            input=synthesis_input, voice=voice, audio_config=audio_config
        )
        
        # Convertimos los bytes binarios del archivo de audio a una cadena en Base64
        audio_base64 = base64.b64encode(response_audio.audio_content).decode("utf-8")
        
        # Retornamos el texto y el string del audio embebido
        return {
            "respuesta": texto_final,
            "audio": audio_base64
        }

    except Exception as e:
        print(f"❌ ERROR EN ENDPOINT /PREGUNTAR: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error interno del servidor: {str(e)}")